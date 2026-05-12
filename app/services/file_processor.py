"""File processor service for reading PDF, DOCX, TXT, CSV, Excel and ZIP files."""

import io
import logging
import os
from typing import Optional

import pandas as pd
import pdfplumber
import pyzipper
from docx import Document

from app.config import APP_CONFIG, settings as app_settings
from app.schemas.internal import FileContent

logger = logging.getLogger(__name__)


class FileProcessingError(Exception):
    """Base exception for file processing errors."""

    def __init__(self, message: str, error_code: str = "INVALID_FILE"):
        self.message = message
        self.error_code = error_code
        super().__init__(message)


class UnsupportedFormatError(FileProcessingError):
    """Raised when a file format is not supported."""

    def __init__(self, filename: str):
        ext = os.path.splitext(filename)[1].lower()
        super().__init__(
            message=f"Unsupported file format: '{ext}' (file: {filename})",
            error_code="UNSUPPORTED_FORMAT",
        )


class InvalidFileError(FileProcessingError):
    """Raised when a file is corrupt or cannot be parsed."""

    def __init__(self, filename: str, reason: str = ""):
        detail = f" — {reason}" if reason else ""
        super().__init__(
            message=f"Cannot read file: '{filename}'{detail}",
            error_code="INVALID_FILE",
        )


class ZipPasswordRequiredError(FileProcessingError):
    """Raised when a ZIP file requires a password but none was provided."""

    def __init__(self, filename: str):
        super().__init__(
            message=f"ZIP file '{filename}' is password-protected. Provide zip_password.",
            error_code="ZIP_PASSWORD_REQUIRED",
        )


class ZipWrongPasswordError(FileProcessingError):
    """Raised when the provided ZIP password is incorrect."""

    def __init__(self, filename: str):
        super().__init__(
            message=f"Wrong password for ZIP file '{filename}'.",
            error_code="ZIP_WRONG_PASSWORD",
        )


class TooManyRowsError(FileProcessingError):
    """Raised when a tabular file exceeds the maximum allowed row count."""

    def __init__(self, filename: str, row_count: int, max_rows: int):
        super().__init__(
            message=f"Dosya '{filename}' çok fazla satır içeriyor: {row_count:,} (limit: {max_rows:,})",
            error_code="TOO_MANY_ROWS",
        )


SUPPORTED_FORMATS = APP_CONFIG["supported_formats"]


class FileProcessor:
    """Reads uploaded files and returns structured FileContent objects."""

    def process(
        self,
        file_bytes: bytes,
        filename: str,
        zip_password: Optional[str] = None,
    ) -> list[FileContent]:
        """Read a file and return a list of FileContent objects.

        For ZIP archives the list may contain multiple entries (one per
        supported file inside the archive).  For all other formats the
        list contains exactly one element.
        """
        ext = os.path.splitext(filename)[1].lower()

        if ext not in SUPPORTED_FORMATS:
            raise UnsupportedFormatError(filename)

        if ext == ".zip":
            return self._extract_zip(file_bytes, zip_password, filename)

        return [self._read_single(file_bytes, filename, ext)]

    # ------------------------------------------------------------------
    # Internal readers
    # ------------------------------------------------------------------

    def _read_single(
        self, data: bytes, filename: str, ext: str
    ) -> FileContent:
        """Dispatch to the correct reader based on extension."""
        if ext == ".pdf":
            text = self._read_pdf(data, filename)
            return FileContent(
                filename=filename,
                content=text,
                file_type="text",
                file_extension=ext,
            )
        if ext == ".docx":
            text = self._read_docx(data, filename)
            return FileContent(
                filename=filename,
                content=text,
                file_type="text",
                file_extension=ext,
            )
        if ext == ".txt":
            text = self._read_txt(data, filename)
            return FileContent(
                filename=filename,
                content=text,
                file_type="text",
                file_extension=ext,
            )
        if ext in {".csv", ".xlsx", ".xls"}:
            df = self._read_tabular(data, filename, ext)
            return FileContent(
                filename=filename,
                content="",
                file_type="table",
                file_extension=ext,
                dataframe=df,
            )

        raise UnsupportedFormatError(filename)

    def _read_pdf(self, data: bytes, filename: str) -> str:
        """Extract text from a PDF using pdfplumber."""
        try:
            pages: list[str] = []
            with pdfplumber.open(io.BytesIO(data)) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        pages.append(text)
            return "\n".join(pages)
        except Exception as exc:
            raise InvalidFileError(filename, reason=str(exc)) from exc

    def _read_docx(self, data: bytes, filename: str) -> str:
        """Extract text from a DOCX file using python-docx.

        Reads both paragraphs and table cells — handles Resmi Gazete
        style documents where all content lives inside tables.
        """
        try:
            doc = Document(io.BytesIO(data))
            parts: list[str] = []

            # 1. Normal paragraflar
            for p in doc.paragraphs:
                if p.text and p.text.strip():
                    parts.append(p.text)

            # 2. Tablo hücreleri — paragraflar boşsa veya tablolarda ek içerik varsa
            seen: set[str] = set(parts)  # duplikasyon önleme
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text and cell_text not in seen:
                            seen.add(cell_text)
                            parts.append(cell_text)

            return "\n".join(parts)
        except Exception as exc:
            raise InvalidFileError(filename, reason=str(exc)) from exc

    def _read_txt(self, data: bytes, filename: str) -> str:
        """Decode a plain-text file trying multiple encodings.

        Tries utf-8, utf-8-sig, windows-1254, iso-8859-9, latin-1,
        iso-8859-1 in order (matching the original Streamlit app).
        """
        encodings = [
            "utf-8",
            "utf-8-sig",
            "windows-1254",
            "iso-8859-9",
            "latin-1",
            "iso-8859-1",
        ]
        for enc in encodings:
            try:
                return data.decode(enc)
            except (UnicodeDecodeError, LookupError):
                continue
        raise InvalidFileError(filename, reason="Unable to decode text with any supported encoding")

    def _read_tabular(
        self, data: bytes, filename: str, ext: str
    ) -> pd.DataFrame:
        """Read CSV or Excel data into a DataFrame."""
        try:
            if ext == ".csv":
                return self._read_csv_with_encoding(data, filename)
            if ext == ".xlsx":
                return self._read_xlsx_fast(data, filename)
            if ext == ".xls":
                df = pd.read_excel(io.BytesIO(data), engine="xlrd")
                if len(df) > app_settings.MAX_TABLE_ROWS:
                    raise TooManyRowsError(filename, len(df), app_settings.MAX_TABLE_ROWS)
                logger.info("xls read: %s — %d satır × %d kolon", filename, len(df), len(df.columns))
                return df
        except (InvalidFileError, TooManyRowsError):
            raise
        except Exception as exc:
            raise InvalidFileError(filename, reason=str(exc)) from exc

        raise UnsupportedFormatError(filename)

    def _read_xlsx_fast(self, data: bytes, filename: str) -> pd.DataFrame:
        """Read xlsx using openpyxl read-only mode for speed."""
        from openpyxl import load_workbook

        try:
            wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
            try:
                ws = wb.active
                logger.info(
                    "xlsx read: %s — ws.max_row=%s, ws.max_column=%s",
                    filename, ws.max_row, ws.max_column,
                )
                rows = []
                header = None
                consecutive_empty = 0
                max_consecutive_empty = 50
                min_filled_cells = 2  # satırda en az 2 kolonda değer olmalı
                max_rows = app_settings.MAX_TABLE_ROWS

                for i, row in enumerate(ws.iter_rows(values_only=True)):
                    if i == 0:
                        header = [str(c) if c is not None else f"col_{j}" for j, c in enumerate(row)]
                        continue

                    # Tamamen boş satır kontrolü
                    if all(c is None or str(c).strip() == "" for c in row):
                        consecutive_empty += 1
                        if consecutive_empty >= max_consecutive_empty:
                            logger.info(
                                "xlsx read: %s — %d ardışık boş satır, okuma durduruluyor (satır %d)",
                                filename, max_consecutive_empty, i,
                            )
                            break
                        continue

                    # Minimum anlamlı hücre filtresi — tek bir "0", " - " vb. içeren
                    # artık satırları ele (consecutive_empty sayacını sıfırlamasın)
                    non_empty = sum(1 for c in row if c is not None and str(c).strip())
                    if non_empty < min_filled_cells:
                        consecutive_empty += 1
                        if consecutive_empty >= max_consecutive_empty:
                            logger.info(
                                "xlsx read: %s — %d ardışık boş/anlamsız satır, okuma durduruluyor (satır %d)",
                                filename, max_consecutive_empty, i,
                            )
                            break
                        continue

                    consecutive_empty = 0
                    rows.append(row)

                    if len(rows) > max_rows:
                        raise TooManyRowsError(filename, len(rows), max_rows)
            finally:
                wb.close()

            if header:
                df = pd.DataFrame(rows, columns=header)
            else:
                df = pd.DataFrame(rows)

            df = df.dropna(how="all").reset_index(drop=True)
            logger.info(
                "xlsx read: %s — %d satır × %d kolon okundu",
                filename, len(df), len(df.columns),
            )
            return df
        except TooManyRowsError:
            raise
        except Exception as exc:
            raise InvalidFileError(filename, reason=str(exc)) from exc


    def _read_csv_with_encoding(self, data: bytes, filename: str) -> pd.DataFrame:
        """Read CSV trying multiple encodings with auto-separator detection."""
        import csv

        # UTF-16 BOM kontrolü — UTF-16 dosyalarda her karakter arasında \x00 olur,
        # NUL temizleme yapısını bozar, bu yüzden önce UTF-8'e dönüştür.
        if data[:2] in (b'\xff\xfe', b'\xfe\xff'):
            try:
                text = data.decode('utf-16')
                data = text.encode('utf-8')
                logger.info("csv UTF-16 BOM algılandı, UTF-8'e dönüştürüldü: %s", filename)
            except (UnicodeDecodeError, LookupError):
                pass

        # NUL byte temizle (bozuk CSV'lerde sniffer'ı yanıltır)
        data = data.replace(b"\x00", b"")

        encodings = ["utf-8", "utf-8-sig", "utf-16", "utf-16-le", "windows-1254", "iso-8859-9", "latin-1"]
        max_rows = app_settings.MAX_TABLE_ROWS
        last_exc = None

        for enc in encodings:
            try:
                text_sample = data[:8192].decode(enc)
            except (UnicodeDecodeError, LookupError):
                continue

            try:
                dialect = csv.Sniffer().sniff(text_sample)
                sep = dialect.delimiter
            except csv.Error:
                # Yaygın delimiter'ları dene
                for candidate_sep in [",", ";", "\t", "|"]:
                    if candidate_sep in text_sample:
                        sep = candidate_sep
                        break
                else:
                    sep = ","

            try:
                df = pd.read_csv(
                    io.BytesIO(data),
                    encoding=enc,
                    sep=sep,
                    engine="c",
                    nrows=max_rows,
                    on_bad_lines="skip",
                )
                df = df.dropna(how="all").reset_index(drop=True)

                if len(df) >= max_rows:
                    raise TooManyRowsError(filename, len(df), max_rows)

                logger.info(
                    "csv read: %s — %d satır × %d kolon okundu (enc=%s, sep=%r)",
                    filename, len(df), len(df.columns), enc, sep,
                )
                return df
            except TooManyRowsError:
                raise
            except (UnicodeDecodeError, Exception) as exc:
                last_exc = exc
                continue

        raise InvalidFileError(filename, reason=f"CSV encoding error: {last_exc}")


    def _extract_zip(
        self,
        data: bytes,
        password: Optional[str],
        zip_filename: str,
        max_decompressed_bytes: int = 500 * 1024 * 1024,  # 500MB limit
    ) -> list[FileContent]:
        """Extract supported files from a ZIP (including AES-encrypted) archive."""
        try:
            with pyzipper.AESZipFile(io.BytesIO(data), "r") as zf:
                # Check if any entry is encrypted
                needs_password = any(
                    info.flag_bits & 0x1 for info in zf.infolist()
                )

                if needs_password and not password:
                    raise ZipPasswordRequiredError(zip_filename)

                pwd_bytes = password.encode("utf-8") if password else None

                # ZIP bomb check: toplam decompressed boyut kontrolü
                total_size = sum(info.file_size for info in zf.infolist() if not info.is_dir())
                if total_size > max_decompressed_bytes:
                    raise InvalidFileError(
                        zip_filename,
                        reason=f"ZIP decompressed boyutu çok büyük: {total_size / (1024*1024):.0f}MB "
                               f"(limit: {max_decompressed_bytes / (1024*1024):.0f}MB)",
                    )

                results: list[FileContent] = []
                for info in zf.infolist():
                    # Skip directories
                    if info.is_dir():
                        continue

                    inner_name = info.filename
                    inner_ext = os.path.splitext(inner_name)[1].lower()

                    if inner_ext not in SUPPORTED_FORMATS or inner_ext == ".zip":
                        # Skip unsupported files inside the ZIP (and nested ZIPs)
                        continue

                    try:
                        member_data = zf.read(inner_name, pwd=pwd_bytes)
                    except RuntimeError as exc:
                        # pyzipper raises RuntimeError for wrong password
                        if "password" in str(exc).lower() or "Bad password" in str(exc):
                            raise ZipWrongPasswordError(zip_filename) from exc
                        raise InvalidFileError(inner_name, reason=str(exc)) from exc

                    results.append(
                        self._read_single(member_data, inner_name, inner_ext)
                    )

                return results

        except (ZipPasswordRequiredError, ZipWrongPasswordError):
            raise
        except FileProcessingError:
            raise
        except Exception as exc:
            # Handle bad password that surfaces as a generic exception
            exc_msg = str(exc).lower()
            if "password" in exc_msg or "bad password" in exc_msg:
                if not password:
                    raise ZipPasswordRequiredError(zip_filename) from exc
                raise ZipWrongPasswordError(zip_filename) from exc
            raise InvalidFileError(zip_filename, reason=str(exc)) from exc
